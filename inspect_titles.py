from docx import Document

files = [
    r"d:\AI\mark123\LGKWFE26Y0955-【无标记稿】一种动水环境衬砌注浆修复模拟装置及模拟方法.docx"
]

doc = Document(files[0])
for i, p in enumerate(doc.paragraphs):
    if i < 20 or i > len(doc.paragraphs) - 30:
        text = p.text.strip()
        if len(text) < 15 and len(text) > 0:
            print(f"[{i}] {text}")

