# -*- coding: utf-8 -*-
"""诊断文档结构：打印文档末尾所有段落"""
from docx import Document
from doc_parser import parse_document
import re

files = [
    r"d:\AI\mark123\FE26Y1586-【无标记初稿】一种新能源汽车专用轴系轴承拆解的夹紧定位工装 - 副本.docx",
    r"d:\AI\mark123\LGKWFE26Y0955-【无标记稿】一种动水环境衬砌注浆修复模拟装置及模拟方法.docx"
]

for f in files:
    doc = Document(f)
    total = len(doc.paragraphs)
    print(f"\n{'='*60}")
    print(f"文件: {f.split(chr(92))[-1]}")
    print(f"总段落数: {total}")
    
    # 打印文档后半部分所有段落（从第70段到末尾）
    start = max(0, total - 80)
    print(f"\n--- 段落 [{start}] ~ [{total-1}] ---")
    for i in range(start, total):
        p = doc.paragraphs[i]
        text = p.text.strip()
        # 标记是否有内容
        has_runs = len(p.runs) > 0
        has_objects = bool(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'))
        has_images = bool(p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip') or
                         p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'))
        
        flags = []
        if has_objects: flags.append("OBJ")
        if has_images: flags.append("IMG")
        if not text and not has_objects and not has_images: flags.append("EMPTY")
        
        flag_str = f" [{','.join(flags)}]" if flags else ""
        
        # 截断过长文本
        display = text[:80] + "..." if len(text) > 80 else text
        print(f"  [{i:3d}]{flag_str} {display}")

    # 打印解析结果
    data = parse_document(f)
    print(f"\n--- 解析出的章节 ---")
    for pos, name in data['title_positions']:
        sec = data['sections'].get(name)
        if sec:
            print(f"  [{pos:3d}] {name} -> 范围 [{sec.start_idx}, {sec.end_idx})")
        else:
            print(f"  [{pos:3d}] {name}")
